"""
Create Word document with integrated Results and Discussion sections.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_paper_revision_doc():
    """Create Word document with integrated sections."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Paper Revision: Integrated Results and Discussion', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    doc.add_paragraph(
        "This document contains the updated Results and Discussion sections ready for "
        "integration into the paper, along with a complete revision guide for remaining fixes."
    )
    
    # RESULTS SECTION
    doc.add_heading('RESULTS SECTION', 1)
    doc.add_paragraph(
        "The following Results section should replace the existing Results section in your paper."
    )
    
    # Read Results content
    with open('RESULTS_SECTION_FINAL.md', 'r', encoding='utf-8') as f:
        results_content = f.read()
    
    # Add Results section (simplified - full text in separate section)
    doc.add_heading('Complete Results Section Text', 2)
    doc.add_paragraph(
        "See 'INTEGRATED_RESULTS_DISCUSSION.md' for the complete Results section text "
        "ready for copy-paste. The section includes:"
    )
    
    results_items = [
        "Evaluation scores table (Table 1) with Finnish, Polish, and Croatian data",
        "Finding 1: Self-scoring bias persists with blind evaluation",
        "Finding 2: Language-specific evaluation patterns",
        "Finding 3: Human translations rated lower",
        "Finding 4: Force Dynamics analysis (100% identification rate)",
        "Comparison with previous findings",
        "Limitations and interpretation",
        "Implications"
    ]
    
    for item in results_items:
        p = doc.add_paragraph(item, style='List Bullet')
    
    # DISCUSSION SECTION
    doc.add_heading('DISCUSSION SECTION', 1)
    doc.add_paragraph(
        "The following Discussion section should replace or update the existing Discussion section."
    )
    
    doc.add_heading('Complete Discussion Section Text', 2)
    doc.add_paragraph(
        "See 'INTEGRATED_RESULTS_DISCUSSION.md' for the complete Discussion section text. "
        "The section includes:"
    )
    
    discussion_items = [
        "Why bias persists with blind evaluation (probabilistic similarity, stylistic alignment)",
        "Language-specific patterns (cross-linguistic variation, morphological complexity)",
        "Human translations and evaluation criteria (literal vs. free preference)",
        "Methodological contributions (blind evaluation, explicit FD framework)",
        "Limitations and future research directions"
    ]
    
    for item in discussion_items:
        p = doc.add_paragraph(item, style='List Bullet')
    
    # REVISION GUIDE
    doc.add_heading('COMPLETE REVISION GUIDE', 1)
    doc.add_paragraph(
        "The following sections provide a complete guide for all remaining revisions needed in the paper."
    )
    
    # Terminological Fixes
    doc.add_heading('1. Terminological Fixes', 2)
    doc.add_paragraph("Apply the following find-and-replace fixes throughout the paper:")
    
    fixes = [
        ("'can process' → 'can analyze'", "When referring to GPT operations"),
        ("'implicit meanings' → 'Force Dynamics relations'", "Be specific about what is meant"),
        ("'AI language processing' → 'linguistic processing by LLMs'", "More accurate terminology"),
        ("'key to advancing' → 'contributes to advancing'", "Tone down overstatement"),
        ("'GPT interprets' → 'GPT analyzes'", "Remove humanizing language"),
        ("'GPT understands' → 'GPT processes'", "Remove humanizing language"),
        ("'demonstrates' → 'suggests'", "Given small sample size"),
        ("'static translation engine' → 'using it solely for direct translation'", "More accurate description")
    ]
    
    for fix, context in fixes:
        p = doc.add_paragraph(f"{fix} ({context})", style='List Bullet')
    
    # Text Replacements
    doc.add_heading('2. Specific Text Replacements', 2)
    doc.add_paragraph(
        "See 'TEXT_REPLACEMENTS_COMPLETE.md' for detailed find-and-replace instructions. "
        "Key replacements include:"
    )
    
    replacements = [
        "Introduction: Fix 'static translation engine' and 'AI language processing'",
        "Related Work: Reorganize NMT vs LLM distinction, fix citations",
        "Methodology: Replace entire section with content from PAPER_METHODOLOGY_SECTION.md",
        "Results: Add explanations for Figure 1 and examples",
        "Discussion: Add limitations section, clarify FD preservation goals"
    ]
    
    for replacement in replacements:
        p = doc.add_paragraph(replacement, style='List Bullet')
    
    # Add Explanations
    doc.add_heading('3. Add Explanations', 2)
    doc.add_paragraph("Add detailed explanations for:")
    
    explanations = [
        "Figure 1: Explain what the figure shows, what bars represent, walk through an example",
        "Examples: For each example, explain FD structure, how each translation preserves/shifts FD, score interpretation",
        "Tables: Explain what data is shown, what patterns are visible"
    ]
    
    for explanation in explanations:
        p = doc.add_paragraph(explanation, style='List Bullet')
    
    # Tone Down Overstatements
    doc.add_heading('4. Tone Down Overstatements', 2)
    doc.add_paragraph("Replace strong claims with more moderate language:")
    
    overstatements = [
        "'key to' → 'contributes to'",
        "'essential' → 'important'",
        "'demonstrates' → 'suggests' (given small sample)",
        "'proves' → 'indicates'",
        "'excel at' → 'show capability in'"
    ]
    
    for overstatement in overstatements:
        p = doc.add_paragraph(overstatement, style='List Bullet')
    
    # Sample Size Caveats
    doc.add_heading('5. Add Sample Size Caveats', 2)
    doc.add_paragraph(
        "Add the following caveat after major claims:"
    )
    caveat = doc.add_paragraph(
        "\"Given the small sample size (10 sentences per language), these findings should be "
        "interpreted as exploratory rather than generalizable.\"",
        style='Quote'
    )
    
    # Key Statistics
    doc.add_heading('KEY STATISTICS FOR INTEGRATION', 1)
    
    stats_table = doc.add_table(rows=4, cols=5)
    stats_table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = stats_table.rows[0].cells
    header_cells[0].text = 'Language'
    header_cells[1].text = 'GPT Self-Score'
    header_cells[2].text = 'Google Translate'
    header_cells[3].text = 'Human Reference'
    header_cells[4].text = 'Bias (Self-Human)'
    
    # Data rows
    data = [
        ['Finnish', '0.940', '0.905', '0.735', '+0.205'],
        ['Polish', '0.935', '0.830', '0.675', '+0.260'],
        ['Croatian', '0.885', '0.930', '0.915', '-0.030']
    ]
    
    for i, row_data in enumerate(data, 1):
        cells = stats_table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            cells[j].text = cell_data
    
    # Force Dynamics Statistics
    doc.add_heading('Force Dynamics Analysis Statistics', 2)
    doc.add_paragraph("Force Dynamics identification rates:")
    
    fd_stats = [
        "Agonist/Antagonist identification: 100% (30/30 per language)",
        "Finnish: Causation (13), Permission (13), Blocking (5)",
        "Polish: Causation (15), Permission (11), Blocking (6)",
        "Croatian: Causation (15), Permission (10), Blocking (6)"
    ]
    
    for stat in fd_stats:
        p = doc.add_paragraph(stat, style='List Bullet')
    
    # Next Steps
    doc.add_heading('NEXT STEPS', 1)
    
    steps = [
        "1. Replace Results section with content from INTEGRATED_RESULTS_DISCUSSION.md",
        "2. Replace/update Discussion section with content from INTEGRATED_RESULTS_DISCUSSION.md",
        "3. Apply terminological fixes using TERMINOLOGY_FIXES.md",
        "4. Apply text replacements using TEXT_REPLACEMENTS_COMPLETE.md",
        "5. Add explanations for Figure 1 and examples using REVISED_SECTIONS.md",
        "6. Tone down overstatements throughout paper",
        "7. Add sample size caveats where appropriate",
        "8. Review Methodology section consistency with Results/Discussion",
        "9. Complete reviewer response document",
        "10. Final proofreading"
    ]
    
    for step in steps:
        p = doc.add_paragraph(step, style='List Number')
    
    # Save document
    doc.save('PAPER_REVISION_INTEGRATED.docx')
    print("Word document created: PAPER_REVISION_INTEGRATED.docx")

if __name__ == '__main__':
    create_paper_revision_doc()
